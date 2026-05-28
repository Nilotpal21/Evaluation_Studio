"""Debug DOCX extraction to see what's being returned"""

import io
import json
import requests
from docx import Document

# Create a DOCX with substantial content
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is the first paragraph with some text content.')
doc.add_paragraph('This is the second paragraph with more text.')

# Add a table
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = 'Header 1'
table.cell(0, 1).text = 'Header 2'
table.cell(1, 0).text = 'Data 1'
table.cell(1, 1).text = 'Data 2'

doc.add_paragraph('This is a third paragraph after the table.')

buffer = io.BytesIO()
doc.save(buffer)
buffer.seek(0)
file_bytes = buffer.read()

print(f"Generated DOCX size: {len(file_bytes)} bytes")

# Extract
files = {'file': ('test.docx', io.BytesIO(file_bytes), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
options = json.dumps({
    'extractImages': True,
    'extractTables': True,
    'preserveLayout': True,
    'renderScreenshots': False,
    'ocrEnabled': False,  # Not needed for DOCX
})
data = {'options': options}

response = requests.post(
    'http://localhost:8080/extract',
    files=files,
    data=data,
    timeout=60,
)

print(f"\nResponse status: {response.status_code}")

if response.status_code == 200:
    result = response.json()
    print(f"\nResponse keys: {list(result.keys())}")
    print(f"Number of pages: {len(result.get('pages', []))}")

    if 'pages' in result and len(result['pages']) > 0:
        print(f"\nFirst page text preview: {result['pages'][0].get('text', '')[:200]}")
        print(f"First page keys: {list(result['pages'][0].keys())}")
    else:
        print("\nNo pages extracted!")

    # Print full response for debugging
    print(f"\n=== FULL RESPONSE ===")
    print(json.dumps(result, indent=2))
else:
    print(f"Error: {response.text}")
