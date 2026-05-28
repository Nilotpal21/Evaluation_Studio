"""
Unified Document Extraction Service

FastAPI microservice for extracting structured content from documents using:
- IBM Docling (PDF, DOCX, PPTX, HTML, images)
- LlamaIndex (TXT, Markdown, JSON, CSV, XML)

Endpoints:
- POST /extract: Extract pages with layout, tables, images, and screenshots
- GET /health: Health check
"""

from __future__ import annotations

import asyncio
import base64
import io
import logging
import tempfile
import os
from typing import Any, Dict, List, Optional
from datetime import datetime
from pathlib import Path as FilePath

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
import uvicorn

# Configure logging (must be before LlamaIndex import)
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


# ─── GPU Detection ──────────────────────────────────────────────────────────
# Detect GPU availability at module load time. This runs once at process startup.
# torch.cuda.is_available() returns True only when an NVIDIA GPU device is
# mounted into the container (NVIDIA Container Toolkit + gpu.enabled=true in Helm).
# On CPU-only hosts, the CUDA libraries in the image are inert.

def detect_device() -> tuple:
    """Detect compute device (CUDA GPU or CPU) at startup."""
    try:
        import torch
        if torch.cuda.is_available():
            device_name = torch.cuda.get_device_name(0)
            logger.info(f"GPU detected: {device_name} (CUDA {torch.version.cuda})")
            return "cuda", device_name
        else:
            logger.info("No GPU detected — running on CPU")
            return "cpu", None
    except ImportError:
        logger.warning("torch not available — running on CPU (torch import failed)")
        return "cpu", None
    except Exception as e:
        logger.warning(f"GPU detection failed — running on CPU: {e}")
        return "cpu", None


DEVICE, GPU_DEVICE_NAME = detect_device()

# Auto-limit workers when GPU is detected (CUDA not fork-safe).
# Helm also sets WORKERS=1 when gpu.enabled=true as a safety net.
if DEVICE == "cuda" and os.environ.get("WORKERS") is None:
    os.environ["WORKERS"] = "1"
    logger.info("GPU detected: auto-setting WORKERS=1 (CUDA is not fork-safe)")

# Import tokenizer for accurate token counting
from tokenizer import count_tokens

# LlamaIndex imports (lazy loaded to avoid startup overhead)
try:
    from llama_index.core import SimpleDirectoryReader, Document as LlamaDocument
    from llama_index.core.node_parser import SentenceSplitter
    from llama_index.readers.file import MarkdownReader, CSVReader, XMLReader

    LLAMAINDEX_AVAILABLE = True
except ImportError:
    logger.warning("LlamaIndex not available - text format extraction will be disabled")
    LLAMAINDEX_AVAILABLE = False

# Language detection imports
try:
    from language_detector import get_language_detector

    LANGUAGE_DETECTION_AVAILABLE = True
except ImportError:
    logger.warning(
        "Language detection not available - install lingua-language-detector and fasttext"
    )
    LANGUAGE_DETECTION_AVAILABLE = False

# Initialize FastAPI
app = FastAPI(
    title="Docling Extraction Service",
    description="Document extraction service using IBM Docling",
    version="1.0.0",
)

# ─── Models ──────────────────────────────────────────────────────────────────


class ExtractionOptions(BaseModel):
    """Options for document extraction"""

    extractImages: bool = Field(default=True, description="Extract images from pages")
    extractTables: bool = Field(default=True, description="Extract tables with structure")
    preserveLayout: bool = Field(default=True, description="Preserve document layout hierarchy")
    renderScreenshots: bool = Field(default=True, description="Render page screenshots")
    ocrEnabled: bool = Field(default=True, description="Enable OCR for scanned documents")


class BoundingBox(BaseModel):
    """Bounding box coordinates"""

    x: float
    y: float
    width: float
    height: float
    page: int


class HeadingInfo(BaseModel):
    """Heading information"""

    level: int
    text: str
    bbox: Optional[BoundingBox] = None


class TableInfo(BaseModel):
    """Extracted table information"""

    rows: List[List[str]]
    headers: List[str]
    html: str
    markdown: str
    bbox: Optional[BoundingBox] = None
    isComplete: bool = True


class ImageInfo(BaseModel):
    """Extracted image information"""

    data: str  # Base64 encoded
    format: str
    bbox: Optional[BoundingBox] = None


class PageData(BaseModel):
    """Single page data"""

    pageNumber: int
    text: str
    layout: Dict[str, Any]
    tables: List[TableInfo]
    images: List[ImageInfo]
    screenshot: Optional[str] = None  # Base64 encoded


class ExtractionMetadata(BaseModel):
    """Extraction metadata"""

    pageCount: int
    hasOCR: bool
    totalTables: int
    totalImages: int
    processingTime: float
    documentType: Optional[str] = None
    # Language detection fields
    language: Optional[str] = Field(default=None, description="Detected language (ISO 639-1 code)")
    languageConfidence: Optional[float] = Field(
        default=None, description="Language detection confidence (0.0-1.0)"
    )
    languageScript: Optional[str] = Field(
        default=None, description="Detected script (Latin, CJK, Arabic, etc.)"
    )
    languageDetectionMethod: Optional[str] = Field(
        default=None, description="Detection method used"
    )
    secondaryLanguages: Optional[List[Dict[str, Any]]] = Field(
        default=None, description="Secondary languages detected"
    )
    # Document metadata fields
    author: Optional[str] = Field(default=None, description="Document author")
    createdDate: Optional[str] = Field(
        default=None, description="Document creation date (ISO 8601)"
    )
    modifiedDate: Optional[str] = Field(
        default=None, description="Document modification date (ISO 8601)"
    )
    title: Optional[str] = Field(default=None, description="Document title")
    subject: Optional[str] = Field(default=None, description="Document subject")
    keywords: Optional[List[str]] = Field(default=None, description="Document keywords")


class DocumentStructure(BaseModel):
    """Document structure information"""

    outline: List[Dict[str, Any]] = []
    documentType: Optional[str] = None


class ExtractionResult(BaseModel):
    """Complete extraction result"""

    pages: List[PageData]
    metadata: ExtractionMetadata
    structure: DocumentStructure


# ─── Format Constants ────────────────────────────────────────────────────────

# MIME types handled by Docling (11 formats - validated 2026-03-26)
# Docling natively supports: PDF, Office docs (modern formats only), HTML, images, Markdown
# Legacy formats (.doc, .ppt) removed - require LibreOffice conversion which is not available
DOCLING_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",  # .pptx
    "text/html",
    "text/markdown",  # Docling supports MD natively
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/tiff",
    "image/bmp",
    "image/webp",
}

# MIME types handled by LlamaIndex (1 format)
# Plain text only - extracted as single page, chunked in downstream processing
# CSV, JSON, XML removed - need specialized hierarchical tree extraction (see task #15)
LLAMAINDEX_TYPES = {
    "text/plain",
}


# ─── Helper Functions ────────────────────────────────────────────────────────


def encode_image_as_png_base64(pil_image: Any) -> str:
    """
    Encode a PIL image as base64 PNG.
    Returns empty string when no image is available.
    """
    if pil_image is None:
        return ""

    buffer = io.BytesIO()
    pil_image.save(buffer, format="PNG")
    buffer.seek(0)
    return base64.b64encode(buffer.read()).decode("utf-8")


def resolve_docling_picture_image(picture: Any, doc: Any | None = None) -> Any:
    """
    Resolve a PIL image from a Docling picture-like object.

    Docling pictures may expose image bytes via an embedded ImageRef or through
    get_image(doc). Older call sites that only probe `pil_image` return empty
    payloads even when the picture data is present.
    """
    if picture is None:
        return None

    direct_image = getattr(picture, "pil_image", None)
    if direct_image is not None:
        return direct_image

    image_ref = getattr(picture, "image", None)
    embedded_image = getattr(image_ref, "pil_image", None) if image_ref is not None else None
    if embedded_image is not None:
        return embedded_image

    if doc is not None:
        get_image = getattr(picture, "get_image", None)
        if callable(get_image):
            try:
                return get_image(doc)
            except Exception as exc:
                logger.warning(f"Error resolving Docling picture image: {exc}")

    return None


def extract_document_metadata(content: bytes, filename: str, content_type: str) -> Dict[str, Any]:
    """
    Extract metadata from document (author, dates, title, etc.).

    Args:
        content: Raw file bytes
        filename: Original filename
        content_type: MIME type

    Returns:
        Dictionary with metadata fields (author, createdDate, modifiedDate, title, subject, keywords)
    """
    metadata = {
        "author": None,
        "createdDate": None,
        "modifiedDate": None,
        "title": None,
        "subject": None,
        "keywords": None,
    }

    try:
        # PDF metadata extraction
        if content_type == "application/pdf":
            try:
                from PyPDF2 import PdfReader

                pdf_file = io.BytesIO(content)
                reader = PdfReader(pdf_file)

                if reader.metadata:
                    # Author
                    if hasattr(reader.metadata, "author") and reader.metadata.author:
                        metadata["author"] = reader.metadata.author
                    elif "/Author" in reader.metadata:
                        metadata["author"] = reader.metadata["/Author"]

                    # Title
                    if hasattr(reader.metadata, "title") and reader.metadata.title:
                        metadata["title"] = reader.metadata.title
                    elif "/Title" in reader.metadata:
                        metadata["title"] = reader.metadata["/Title"]

                    # Subject
                    if hasattr(reader.metadata, "subject") and reader.metadata.subject:
                        metadata["subject"] = reader.metadata.subject
                    elif "/Subject" in reader.metadata:
                        metadata["subject"] = reader.metadata["/Subject"]

                    # Creation date
                    if hasattr(reader.metadata, "creation_date") and reader.metadata.creation_date:
                        metadata["createdDate"] = reader.metadata.creation_date.isoformat()
                    elif "/CreationDate" in reader.metadata:
                        # Parse PDF date format (D:YYYYMMDDHHmmSS)
                        date_str = reader.metadata["/CreationDate"]
                        if date_str and date_str.startswith("D:"):
                            try:
                                from dateutil import parser

                                # Remove D: prefix and timezone
                                clean_date = date_str[2:16]  # YYYYMMDDHHmmSS
                                dt = datetime.strptime(clean_date, "%Y%m%d%H%M%S")
                                metadata["createdDate"] = dt.isoformat()
                            except Exception:
                                pass

                    # Modification date
                    if (
                        hasattr(reader.metadata, "modification_date")
                        and reader.metadata.modification_date
                    ):
                        metadata["modifiedDate"] = reader.metadata.modification_date.isoformat()
                    elif "/ModDate" in reader.metadata:
                        date_str = reader.metadata["/ModDate"]
                        if date_str and date_str.startswith("D:"):
                            try:
                                clean_date = date_str[2:16]
                                dt = datetime.strptime(clean_date, "%Y%m%d%H%M%S")
                                metadata["modifiedDate"] = dt.isoformat()
                            except Exception:
                                pass

                    # Keywords
                    if hasattr(reader.metadata, "keywords") and reader.metadata.keywords:
                        keywords_str = reader.metadata.keywords
                        metadata["keywords"] = [
                            k.strip() for k in keywords_str.split(",") if k.strip()
                        ]
                    elif "/Keywords" in reader.metadata:
                        keywords_str = reader.metadata["/Keywords"]
                        if keywords_str:
                            metadata["keywords"] = [
                                k.strip() for k in keywords_str.split(",") if k.strip()
                            ]

                logger.info(
                    f"Extracted PDF metadata: author={metadata['author']}, title={metadata['title']}"
                )

            except ImportError:
                logger.warning("PyPDF2 not available - skipping PDF metadata extraction")
            except Exception as e:
                logger.warning(f"Failed to extract PDF metadata: {e}")

        # DOCX metadata extraction
        elif (
            content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            try:
                from docx import Document as DocxDocument

                # Save to temp file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name

                try:
                    doc = DocxDocument(tmp_path)
                    core_props = doc.core_properties

                    if core_props:
                        # Author
                        if core_props.author:
                            metadata["author"] = core_props.author

                        # Title
                        if core_props.title:
                            metadata["title"] = core_props.title

                        # Subject
                        if core_props.subject:
                            metadata["subject"] = core_props.subject

                        # Keywords
                        if core_props.keywords:
                            metadata["keywords"] = [
                                k.strip() for k in core_props.keywords.split(",") if k.strip()
                            ]

                        # Created date
                        if core_props.created:
                            metadata["createdDate"] = core_props.created.isoformat()

                        # Modified date
                        if core_props.modified:
                            metadata["modifiedDate"] = core_props.modified.isoformat()

                    logger.info(
                        f"Extracted DOCX metadata: author={metadata['author']}, title={metadata['title']}"
                    )

                finally:
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass

            except ImportError:
                logger.warning("python-docx not available - skipping DOCX metadata extraction")
            except Exception as e:
                logger.warning(f"Failed to extract DOCX metadata: {e}")

        # PPTX metadata extraction
        elif (
            content_type
            == "application/vnd.openxmlformats-officedocument.presentationml.presentation"
        ):
            try:
                from pptx import Presentation

                # Save to temp file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pptx") as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name

                try:
                    prs = Presentation(tmp_path)
                    core_props = prs.core_properties

                    if core_props:
                        # Author
                        if core_props.author:
                            metadata["author"] = core_props.author

                        # Title
                        if core_props.title:
                            metadata["title"] = core_props.title

                        # Subject
                        if core_props.subject:
                            metadata["subject"] = core_props.subject

                        # Keywords
                        if core_props.keywords:
                            metadata["keywords"] = [
                                k.strip() for k in core_props.keywords.split(",") if k.strip()
                            ]

                        # Created date
                        if core_props.created:
                            metadata["createdDate"] = core_props.created.isoformat()

                        # Modified date
                        if core_props.modified:
                            metadata["modifiedDate"] = core_props.modified.isoformat()

                    logger.info(
                        f"Extracted PPTX metadata: author={metadata['author']}, title={metadata['title']}"
                    )

                finally:
                    try:
                        os.unlink(tmp_path)
                    except Exception:
                        pass

            except ImportError:
                logger.warning("python-pptx not available - skipping PPTX metadata extraction")
            except Exception as e:
                logger.warning(f"Failed to extract PPTX metadata: {e}")

        else:
            logger.debug(f"No metadata extraction implemented for content type: {content_type}")

    except Exception as e:
        logger.error(f"Error extracting document metadata: {e}", exc_info=True)

    return metadata


def extract_heading_hierarchy(page: Any) -> List[HeadingInfo]:
    """Extract heading hierarchy from a Docling page"""
    headings = []

    try:
        # Docling provides structured layout elements
        # We'll extract headings from the layout tree
        if hasattr(page, "elements"):
            for element in page.elements:
                if hasattr(element, "label") and "heading" in element.label.lower():
                    # Extract heading level (e.g., "heading-1" -> 1)
                    level = 1
                    if "-" in element.label:
                        try:
                            level = int(element.label.split("-")[1])
                        except (ValueError, IndexError):
                            level = 1

                    heading = HeadingInfo(
                        level=level,
                        text=element.text if hasattr(element, "text") else "",
                        bbox=extract_bbox(element) if hasattr(element, "bbox") else None,
                    )
                    headings.append(heading)
    except Exception as e:
        logger.warning(f"Error extracting headings: {e}")

    return headings


def extract_bbox(element: Any) -> Optional[BoundingBox]:
    """Extract bounding box from element"""
    try:
        if hasattr(element, "bbox") and element.bbox:
            bbox = element.bbox
            return BoundingBox(
                x=bbox.l if hasattr(bbox, "l") else bbox.x if hasattr(bbox, "x") else 0,
                y=bbox.t if hasattr(bbox, "t") else bbox.y if hasattr(bbox, "y") else 0,
                width=bbox.r - bbox.l
                if hasattr(bbox, "r") and hasattr(bbox, "l")
                else bbox.width
                if hasattr(bbox, "width")
                else 0,
                height=bbox.b - bbox.t
                if hasattr(bbox, "b") and hasattr(bbox, "t")
                else bbox.height
                if hasattr(bbox, "height")
                else 0,
                page=getattr(element, "page", 0),
            )
    except Exception as e:
        logger.warning(f"Error extracting bbox: {e}")
    return None


def should_include_image(
    bbox: Optional[BoundingBox], page_width: float = 612.0, page_height: float = 792.0
) -> bool:
    """
    Filter out logos, icons, and decorative images based on size and position heuristics.

    Standard US Letter page dimensions: 612x792 points (8.5x11 inches at 72 DPI)

    Filtering criteria:
    1. Dimension filter: Logos/icons typically < 150x150 points
    2. Position filter: Skip headers (top 10%), footers (bottom 10%), margins (left/right 10%)
    3. Aspect ratio filter: Very wide (>5:1) or tall (<1:5) = decorative borders
    4. Minimum area: Less than 15,000 sq points (~122x122px) likely decorative

    Args:
        bbox: Image bounding box (can be None if position unknown)
        page_width: Page width in points (default: US Letter 612pt)
        page_height: Page height in points (default: US Letter 792pt)

    Returns:
        True if image should be included, False if likely logo/decorative
    """
    # If no bbox, include by default (can't filter unknown position)
    if not bbox:
        return True

    try:
        # 1. Dimension filter: Logos/icons are typically small
        # Common logo sizes: 50x50, 100x100, 150x50, etc.
        if bbox.width < 150 and bbox.height < 150:
            logger.debug(f"Filtered small image: {bbox.width}x{bbox.height}pt (likely logo/icon)")
            return False

        # 2. Position filter: Skip header/footer/margin regions
        # Header (top 10% of page)
        if bbox.y < page_height * 0.1:
            logger.debug(
                f"Filtered header image at y={bbox.y}pt (top {(bbox.y / page_height) * 100:.1f}%)"
            )
            return False

        # Footer (bottom 10% of page)
        if bbox.y + bbox.height > page_height * 0.9:
            logger.debug(
                f"Filtered footer image at y={bbox.y + bbox.height}pt (bottom {((page_height - bbox.y - bbox.height) / page_height) * 100:.1f}%)"
            )
            return False

        # Left margin (left 10% of page)
        if bbox.x < page_width * 0.1:
            logger.debug(
                f"Filtered left margin image at x={bbox.x}pt (left {(bbox.x / page_width) * 100:.1f}%)"
            )
            return False

        # Right margin (right 10% of page)
        if bbox.x + bbox.width > page_width * 0.9:
            logger.debug(
                f"Filtered right margin image at x={bbox.x + bbox.width}pt (right {((page_width - bbox.x - bbox.width) / page_width) * 100:.1f}%)"
            )
            return False

        # 3. Aspect ratio filter: Very wide or tall images are likely decorative
        aspect_ratio = bbox.width / bbox.height if bbox.height > 0 else 0
        if aspect_ratio > 5.0:  # Very wide (e.g., horizontal dividers)
            logger.debug(f"Filtered wide decorative image: aspect ratio {aspect_ratio:.2f}:1")
            return False
        if aspect_ratio < 0.2:  # Very tall (e.g., vertical dividers)
            logger.debug(f"Filtered tall decorative image: aspect ratio 1:{1 / aspect_ratio:.2f}")
            return False

        # 4. Minimum area threshold: Very small images likely decorative
        area = bbox.width * bbox.height
        if area < 15000:  # ~122x122 points
            logger.debug(
                f"Filtered small area image: {area:.0f} sq pt (~{area**0.5:.0f}x{area**0.5:.0f}pt)"
            )
            return False

        # Passed all filters - include this image
        return True

    except Exception as e:
        logger.warning(f"Error in image filtering: {e}")
        # On error, include by default (safe fallback)
        return True


def safe_table_rows(data: Any) -> List[List[str]]:
    """
    Convert Docling table data to List[List[str]].

    Docling's table.data can contain mixed types:
    - str values (simple cells)
    - int/float values (numeric cells)
    - TableCell objects with .text attribute
    - Nested lists of TableCell objects (merged cells)

    This function normalizes everything to strings.
    """
    if not data:
        return []
    rows: List[List[str]] = []
    for row in data:
        if not isinstance(row, (list, tuple)):
            continue
        str_row: List[str] = []
        for cell in row:
            if isinstance(cell, str):
                str_row.append(cell)
            elif isinstance(cell, (int, float)):
                str_row.append(str(cell))
            elif hasattr(cell, "text"):
                # TableCell object
                str_row.append(str(cell.text) if cell.text is not None else "")
            elif isinstance(cell, (list, tuple)):
                # Merged cells: list of TableCell objects — join their texts
                parts = []
                for sub in cell:
                    if hasattr(sub, "text"):
                        parts.append(str(sub.text) if sub.text is not None else "")
                    else:
                        parts.append(str(sub))
                str_row.append(" ".join(parts))
            else:
                str_row.append(str(cell))
        rows.append(str_row)
    return rows


def safe_table_headers(headers: Any) -> List[str]:
    """
    Convert Docling table headers to List[str].
    Handles the same mixed types as safe_table_rows for individual cells.
    """
    if not headers:
        return []
    result: List[str] = []
    for h in headers:
        if isinstance(h, str):
            result.append(h)
        elif isinstance(h, (int, float)):
            result.append(str(h))
        elif hasattr(h, "text"):
            result.append(str(h.text) if h.text is not None else "")
        elif isinstance(h, (list, tuple)):
            parts = []
            for sub in h:
                if hasattr(sub, "text"):
                    parts.append(str(sub.text) if sub.text is not None else "")
                else:
                    parts.append(str(sub))
            result.append(" ".join(parts))
        else:
            result.append(str(h))
    return result


def detect_table_completeness(table: Any, next_page: Optional[Any]) -> bool:
    """
    Detect if a table continues to the next page.
    A table is incomplete if it ends at the page boundary without a clear closing.
    """
    try:
        # Simple heuristic: if table bbox is at the bottom of page, it might continue
        if hasattr(table, "bbox"):
            bbox = table.bbox
            # If table extends to within 50 points of page bottom, mark as potentially incomplete
            # This is a heuristic - you can make it more sophisticated
            if hasattr(bbox, "b") and hasattr(bbox, "page_height"):
                if bbox.page_height - bbox.b < 50:
                    return False

        return True
    except Exception as e:
        logger.warning(f"Error detecting table completeness: {e}")
        return True


def render_page_screenshot(page: Any, pdf_document: Any) -> Optional[str]:
    """
    Render a page as PNG screenshot.
    Returns base64 encoded PNG.
    """
    try:
        # This is a placeholder - actual implementation depends on Docling's API
        # Docling might provide page.render() or similar method

        # For now, we'll use a simple approach with pdf2image if available
        try:
            from pdf2image import convert_from_bytes
            import PIL.Image

            # Get page number
            page_num = getattr(page, "page_no", 0)

            # Convert page to image
            # Note: This is simplified - in production you'd cache the PDF and render pages efficiently
            images = convert_from_bytes(
                pdf_document,
                first_page=page_num + 1,
                last_page=page_num + 1,
                dpi=150,  # Lower DPI for faster processing
            )

            if images:
                img = images[0]
                return encode_image_as_png_base64(img)

        except ImportError:
            logger.warning("pdf2image not available, skipping screenshot rendering")
            return None

    except Exception as e:
        logger.error(f"Error rendering page screenshot: {e}")
        return None


def load_text_document(file_path: str, content_type: str) -> List[LlamaDocument]:
    """
    Load text-based documents using LlamaIndex readers.

    Args:
        file_path: Path to the document file
        content_type: MIME type of the document

    Returns:
        List of LlamaIndex Document objects
    """
    if not LLAMAINDEX_AVAILABLE:
        raise RuntimeError("LlamaIndex not available - cannot load text documents")

    logger.info(f"Loading text document with LlamaIndex: {content_type}")

    try:
        # Route to appropriate loader based on content type
        if content_type == "text/plain":
            # Plain text - use SimpleDirectoryReader
            reader = SimpleDirectoryReader(input_files=[file_path])
            documents = reader.load_data()

        elif content_type == "text/markdown":
            # Markdown - use MarkdownReader for structure awareness
            reader = MarkdownReader()
            documents = reader.load_data(FilePath(file_path))

        elif content_type == "application/json":
            # JSON - use SimpleDirectoryReader (reads as text)
            reader = SimpleDirectoryReader(input_files=[file_path])
            documents = reader.load_data()

        elif content_type == "text/csv":
            # CSV - use CSVReader (needs Path object)
            reader = CSVReader()
            documents = reader.load_data(FilePath(file_path))

        elif content_type in ["application/xml", "text/xml"]:
            # XML - use XMLReader (needs Path object)
            reader = XMLReader()
            documents = reader.load_data(FilePath(file_path))

        else:
            raise ValueError(f"Unsupported text content type: {content_type}")

        logger.info(f"Loaded {len(documents)} document(s) from {content_type}")

        # Add metadata
        for doc in documents:
            if not hasattr(doc, "metadata") or doc.metadata is None:
                doc.metadata = {}
            doc.metadata["content_type"] = content_type
            doc.metadata["source_file"] = file_path

        return documents

    except Exception as e:
        logger.error(f"Error loading text document: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, detail=f"Failed to load {content_type} document: {str(e)}"
        )


def convert_text_to_single_page(
    documents: List[LlamaDocument], options: ExtractionOptions
) -> List[PageData]:
    """
    Convert plain text documents to a single PageData object.

    NOTE: We do NOT chunk here - chunking happens downstream in page-processing-worker
    to avoid double chunking. This maintains consistency with other formats (PDF, DOCX)
    which are extracted as pages and then chunked in the pipeline.

    Args:
        documents: List of LlamaIndex Document objects (usually 1 for plain text)
        options: Extraction options (not used for text)

    Returns:
        List with single PageData object containing full text
    """
    logger.info(
        f"Converting {len(documents)} text document(s) to single page (chunking happens downstream)"
    )

    # Concatenate all documents into single text
    full_text = "\n\n".join(doc.text for doc in documents)

    logger.info(
        f"Created single page with {len(full_text)} characters (will be chunked in page-processing)"
    )

    # Return as single page - downstream worker will chunk with consistent strategy
    page_data = PageData(
        pageNumber=1,
        text=full_text,
        layout={
            "headings": [],  # Plain text has no structural headings
            "structure": {},
        },
        tables=[],  # No table extraction for plain text
        images=[],  # No image extraction for plain text
        screenshot=None,  # No screenshots for text documents
    )

    return [page_data]


def detect_document_type(result: Any) -> str:
    """Detect document type based on content"""
    # Simple heuristic based on structure
    # You can make this more sophisticated

    try:
        if hasattr(result, "metadata"):
            if hasattr(result.metadata, "format"):
                return result.metadata.format

        # Default
        return "unknown"
    except Exception:
        return "unknown"


def process_pages(
    result: Any, options: ExtractionOptions, pdf_bytes: Optional[bytes] = None, filename: Optional[str] = None
) -> List[PageData]:
    """Process Docling result into structured page data"""
    pages = []

    try:
        # Check if result has page-based structure (PDFs) or document-based structure (DOCX/PPTX/HTML)
        if result.pages and len(result.pages) > 0:
            # PDF-style: iterate through pages
            # NOTE: result.pages[] are raw Page objects without export_to_markdown(),
            # tables, or images. All content access must go through result.document
            # which provides per-page export via page_no parameter and provenance-based
            # filtering for tables/images.
            doc = result.document if hasattr(result, "document") and result.document else None

            # Determine if this is a standalone image file (already 1-based page_no)
            _IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.webp', '.gif'}
            is_image_file = False
            if filename:
                is_image_file = FilePath(filename).suffix.lower() in _IMAGE_EXTS

            for page in result.pages:
                page_num = getattr(page, "page_no", len(pages))
                # Page numbering rules:
                # - Standalone images: Docling returns page_no=1 (already 1-based)
                # - PDFs: Docling returns page_no that varies by device:
                #     CPU pipeline: page_no is 0-based (0, 1, 2, ...)
                #     GPU pipeline: page_no is 1-based (1, 2, 3, ...)
                # We resolve by checking which key actually exists in doc.pages.
                if is_image_file:
                    # Images always have 1-based page_no from Docling
                    doc_page_no = page_num if page_num >= 1 else 1
                elif doc and hasattr(doc, 'pages') and doc.pages and hasattr(doc.pages, 'keys'):
                    # Resolve page_no against actual doc.pages keys (handles both
                    # CPU 0-based and GPU 1-based page numbering)
                    doc_pages_keys = sorted(doc.pages.keys())
                    if page_num in doc_pages_keys:
                        # page_num is already a valid key (GPU: 1-based match)
                        doc_page_no = page_num
                    elif (page_num + 1) in doc_pages_keys:
                        # page_num is 0-based, doc.pages is 1-based (CPU path)
                        doc_page_no = page_num + 1
                    else:
                        # Neither matches — use first available key for this
                        # iteration index as best-effort fallback
                        page_idx = len(pages)  # 0-based index of pages processed so far
                        if page_idx < len(doc_pages_keys):
                            doc_page_no = doc_pages_keys[page_idx]
                        else:
                            doc_page_no = page_num + 1
                        logger.warning(
                            f"Page numbering mismatch: page_num={page_num} not found "
                            f"in doc.pages keys={doc_pages_keys}, using doc_page_no={doc_page_no}"
                        )
                else:
                    # No doc.pages available — assume 0-based → 1-based
                    doc_page_no = page_num + 1

                # Extract text as Markdown via document-level per-page export
                if doc and hasattr(doc, "export_to_markdown"):
                    page_text = doc.export_to_markdown(page_no=doc_page_no)
                else:
                    # Fallback: try assembled elements text
                    if hasattr(page, "assembled") and hasattr(page.assembled, "elements"):
                        page_text = "\n\n".join(
                            getattr(el, "text", "") for el in page.assembled.elements
                            if getattr(el, "text", "")
                        )
                    else:
                        page_text = ""

                if not page_text or not page_text.strip():
                    logger.info(f"Page {doc_page_no} has no text, skipping")
                    continue

                # Extract headings for this page from document
                headings = []
                if doc:
                    try:
                        for text_item in doc.texts:
                            if (
                                hasattr(text_item, "prov")
                                and text_item.prov
                                and text_item.prov[0].page_no == doc_page_no
                                and hasattr(text_item, "label")
                                and "heading" in str(text_item.label).lower()
                            ):
                                level = 1
                                label_str = str(text_item.label)
                                if "-" in label_str:
                                    try:
                                        level = int(label_str.split("-")[1])
                                    except (ValueError, IndexError):
                                        level = 1
                                headings.append(
                                    HeadingInfo(
                                        level=level,
                                        text=text_item.text if hasattr(text_item, "text") else "",
                                        bbox=extract_bbox(text_item),
                                    )
                                )
                    except Exception as e:
                        logger.warning(f"Error extracting headings for page {doc_page_no}: {e}")
                else:
                    headings = extract_heading_hierarchy(page)

                # Extract tables for this page from document (filtered by provenance)
                tables = []
                if options.extractTables and doc and hasattr(doc, "tables"):
                    for table in doc.tables:
                        try:
                            if (
                                hasattr(table, "prov")
                                and table.prov
                                and table.prov[0].page_no == doc_page_no
                            ):
                                table_info = TableInfo(
                                    rows=safe_table_rows(table.data) if hasattr(table, "data") else [],
                                    headers=safe_table_headers(table.headers) if hasattr(table, "headers") else [],
                                    html=table.export_to_html(doc=doc) if hasattr(table, "export_to_html") else "",
                                    markdown=table.export_to_markdown(doc=doc)
                                    if hasattr(table, "export_to_markdown")
                                    else "",
                                    bbox=extract_bbox(table),
                                    isComplete=True,
                                )
                                tables.append(table_info)
                        except Exception as e:
                            logger.warning(f"Error processing table on page {doc_page_no}: {e}")

                # Extract images for this page from document (filtered by provenance)
                images = []
                if options.extractImages and doc and hasattr(doc, "pictures"):
                    for img in doc.pictures:
                        try:
                            if (
                                hasattr(img, "prov")
                                and img.prov
                                and img.prov[0].page_no == doc_page_no
                            ):
                                bbox = extract_bbox(img)
                                if not should_include_image(bbox):
                                    logger.info(
                                        f"Filtered decorative image on page {doc_page_no}"
                                    )
                                    continue
                                img_data = encode_image_as_png_base64(
                                    resolve_docling_picture_image(img, doc)
                                )
                                if img_data:
                                    images.append(
                                        ImageInfo(data=img_data, format="png", bbox=bbox)
                                    )
                        except Exception as e:
                            logger.warning(f"Error processing image on page {doc_page_no}: {e}")

                # Render screenshot
                screenshot = None
                if options.renderScreenshots and pdf_bytes:
                    screenshot = render_page_screenshot(page, pdf_bytes)

                # Build page data
                page_data = PageData(
                    pageNumber=doc_page_no,
                    text=page_text,
                    layout={"headings": [h.model_dump() for h in headings], "structure": {}},
                    tables=tables,
                    images=images,
                    screenshot=screenshot,
                )

                pages.append(page_data)

        elif hasattr(result, "document") and result.document:
            # Document-style: extract from document object (DOCX/PPTX/HTML)
            doc = result.document

            # Get page numbers from the document's pages dict
            page_numbers = sorted(doc.pages.keys()) if hasattr(doc, "pages") and doc.pages else []

            if len(page_numbers) > 1 and hasattr(doc, "export_to_markdown"):
                # Multi-page document: extract per-page content using page_no filter
                logger.info(f"Document has {len(page_numbers)} pages, extracting per-page")

                for page_num in page_numbers:
                    # Export markdown for this specific page only
                    page_text = doc.export_to_markdown(page_no=page_num)

                    if not page_text or not page_text.strip():
                        logger.info(f"Page {page_num} has no text, skipping")
                        continue

                    # Extract headings for this page
                    page_headings = []
                    try:
                        for text_item in doc.texts:
                            if (
                                hasattr(text_item, "prov")
                                and text_item.prov
                                and text_item.prov[0].page_no == page_num
                                and hasattr(text_item, "label")
                                and "heading" in text_item.label.lower()
                            ):
                                level = 1
                                if "-" in text_item.label:
                                    try:
                                        level = int(text_item.label.split("-")[1])
                                    except (ValueError, IndexError):
                                        level = 1
                                page_headings.append(
                                    HeadingInfo(
                                        level=level,
                                        text=text_item.text if hasattr(text_item, "text") else "",
                                        bbox=extract_bbox(text_item),
                                    )
                                )
                    except Exception as e:
                        logger.warning(f"Error extracting headings for page {page_num}: {e}")

                    # Extract tables for this page
                    page_tables = []
                    if options.extractTables and hasattr(doc, "tables"):
                        for table in doc.tables:
                            try:
                                if (
                                    hasattr(table, "prov")
                                    and table.prov
                                    and table.prov[0].page_no == page_num
                                ):
                                    table_info = TableInfo(
                                        rows=table.data if hasattr(table, "data") else [],
                                        headers=table.headers if hasattr(table, "headers") else [],
                                        html=table.export_to_html() if hasattr(table, "export_to_html") else "",
                                        markdown=table.export_to_markdown()
                                        if hasattr(table, "export_to_markdown")
                                        else "",
                                        bbox=extract_bbox(table),
                                        isComplete=True,
                                    )
                                    page_tables.append(table_info)
                            except Exception as e:
                                logger.warning(f"Error processing table on page {page_num}: {e}")

                    # Extract images for this page with logo/decorative filtering
                    page_images = []
                    if options.extractImages and hasattr(doc, "pictures"):
                        for img in doc.pictures:
                            try:
                                if (
                                    hasattr(img, "prov")
                                    and img.prov
                                    and img.prov[0].page_no == page_num
                                ):
                                    bbox = extract_bbox(img)
                                    if not should_include_image(bbox):
                                        continue
                                    img_data = encode_image_as_png_base64(
                                        resolve_docling_picture_image(img, doc)
                                    )
                                    if img_data:
                                        page_images.append(
                                            ImageInfo(data=img_data, format="png", bbox=bbox)
                                        )
                            except Exception as e:
                                logger.warning(f"Error processing image on page {page_num}: {e}")

                    page_data = PageData(
                        pageNumber=page_num,  # Already 1-based from Docling
                        text=page_text,
                        layout={"headings": [h.model_dump() for h in page_headings], "structure": {}},
                        tables=page_tables,
                        images=page_images,
                        screenshot=None,
                    )
                    pages.append(page_data)

                logger.info(f"Extracted {len(pages)} pages from multi-page document")

            else:
                # Single-page document or no pages dict: fallback to whole-document extraction
                page_text = doc.export_to_markdown() if hasattr(doc, "export_to_markdown") else str(doc)

                headings = extract_heading_hierarchy(doc)

                tables = []
                if options.extractTables and hasattr(doc, "tables"):
                    for table in doc.tables:
                        try:
                            table_info = TableInfo(
                                rows=table.data if hasattr(table, "data") else [],
                                headers=table.headers if hasattr(table, "headers") else [],
                                html=table.export_to_html() if hasattr(table, "export_to_html") else "",
                                markdown=table.export_to_markdown()
                                if hasattr(table, "export_to_markdown")
                                else "",
                                bbox=extract_bbox(table),
                                isComplete=True,
                            )
                            tables.append(table_info)
                        except Exception as e:
                            logger.warning(f"Error processing table: {e}")

                images = []
                if options.extractImages and hasattr(doc, "pictures"):
                    for img in doc.pictures:
                        try:
                            bbox = extract_bbox(img)
                            if not should_include_image(bbox):
                                continue
                            img_data = encode_image_as_png_base64(
                                resolve_docling_picture_image(img, doc)
                            )
                            if img_data:
                                images.append(ImageInfo(data=img_data, format="png", bbox=bbox))
                        except Exception as e:
                            logger.warning(f"Error processing image: {e}")

                page_data = PageData(
                    pageNumber=1,
                    text=page_text,
                    layout={"headings": [h.model_dump() for h in headings], "structure": {}},
                    tables=tables,
                    images=images,
                    screenshot=None,
                )
                pages.append(page_data)

    except Exception as e:
        logger.error(f"Error processing pages: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing document pages: {str(e)}")

    return pages


async def extract_with_docling(
    content: bytes, filename: str, opts: ExtractionOptions
) -> tuple[List[PageData], str]:
    """
    Extract document using Docling engine.

    Args:
        content: Raw file bytes
        filename: Original filename
        opts: Extraction options

    Returns:
        Tuple of (pages, document_type)
    """
    try:
        from docling.document_converter import DocumentConverter, PdfFormatOption, ImageFormatOption
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import (
            PdfPipelineOptions,
            RapidOcrOptions,
            TableFormerMode,
            AcceleratorOptions,
            AcceleratorDevice,
        )
        from docling.document_converter import ConversionResult

        logger.info(f"Extracting with Docling: {filename} (device={DEVICE})")

        # Save to temp file (Docling requires file path)
        suffix = FilePath(filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            # Configure pipeline with explicit table detection
            # TableFormerMode.ACCURATE enables full table structure detection (headers, rows, cells)
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_table_structure = opts.extractTables  # Enable table structure detection
            pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE if opts.extractTables else TableFormerMode.FAST
            pipeline_options.do_ocr = opts.ocrEnabled  # Enable OCR if requested
            if opts.ocrEnabled:
                # For PDFs: only OCR image regions (text-layer extraction handles the rest)
                # force_full_page_ocr=True is reserved for standalone images where the entire file IS the page
                pipeline_options.ocr_options = RapidOcrOptions(force_full_page_ocr=False)
            # Picture extraction requires Docling to materialize picture/page images;
            # otherwise picture items exist but serialize as empty payloads.
            pipeline_options.generate_picture_images = opts.extractImages
            pipeline_options.generate_page_images = opts.extractImages
            pipeline_options.images_scale = 2.0 if opts.extractImages else 1.0

            # GPU acceleration: AcceleratorOptions tells docling's TableFormer and
            # layout models to use CUDA when available. AUTO mode picks
            # CUDA → MPS → CPU in order. num_threads controls parallelism.
            pipeline_options.accelerator_options = AcceleratorOptions(
                device=AcceleratorDevice.AUTO,
                num_threads=int(os.environ.get("DOCLING_NUM_THREADS", "4")),
            )

            # Initialize Docling converter with table detection enabled
            # For images, force full-page OCR since the image IS the page
            IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.webp', '.gif'}
            file_suffix = FilePath(filename).suffix.lower()
            # Detect image by extension; if no extension, use mimetypes guess from filename
            is_image = file_suffix in IMAGE_EXTENSIONS
            if not is_image and not file_suffix:
                import mimetypes
                guessed_type, _ = mimetypes.guess_type(filename)
                if guessed_type and guessed_type.startswith('image/'):
                    is_image = True
                    logger.info(f"Detected image via mime guess: {guessed_type}")
            if is_image and opts.ocrEnabled:
                image_pipeline_options = PdfPipelineOptions()
                image_pipeline_options.do_ocr = True
                image_pipeline_options.ocr_options = RapidOcrOptions(force_full_page_ocr=True)
                image_pipeline_options.do_table_structure = opts.extractTables
                image_pipeline_options.generate_picture_images = opts.extractImages
                image_pipeline_options.generate_page_images = opts.extractImages
                image_pipeline_options.images_scale = 2.0 if opts.extractImages else 1.0
                image_format_opt = ImageFormatOption(pipeline_options=image_pipeline_options)
            else:
                image_format_opt = ImageFormatOption(pipeline_options=pipeline_options)

            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options),
                    InputFormat.IMAGE: image_format_opt,
                }
            )

            # Convert document — run in thread pool to avoid blocking the event loop.
            # converter.convert() is synchronous and can take 5-60s on large files.
            # Without to_thread(), the async event loop freezes and health check
            # probes timeout → K8s restarts the pod thinking it's dead.
            result: ConversionResult = await asyncio.to_thread(converter.convert, tmp_path)

            try:
                import torch
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
            except Exception:
                pass

            # Process pages using existing helper
            pages = process_pages(
                result, opts, pdf_bytes=content if filename.lower().endswith(".pdf") else None, filename=filename
            )

            # ── Standalone image fix ──────────────────────────────────────────
            # When the uploaded file IS an image (PNG, JPEG, etc.), Docling treats
            # the whole file as a single page and runs OCR.  However doc.pictures
            # (embedded images *within* a document) is empty because the image is
            # the document itself, not something embedded inside it.
            # Downstream workers (visual-enrichment, multimodal) rely on
            # page.images to trigger LLM vision analysis, so we inject the
            # original uploaded image into the first page's images array.
            # Reuse the is_image flag computed above (extension + mime fallback)
            if is_image:
                original_image_b64 = base64.b64encode(content).decode("utf-8")
                img_fmt = file_suffix.lstrip('.') if file_suffix else 'png'
                # Normalize short extensions
                img_fmt = {'jpg': 'jpeg', 'tif': 'tiff'}.get(img_fmt, img_fmt)

                original_image = ImageInfo(data=original_image_b64, format=img_fmt, bbox=None)

                if not pages:
                    # OCR found no readable text — create a synthetic page
                    pages = [PageData(
                        pageNumber=1,
                        text="",
                        layout={"headings": [], "structure": {}},
                        tables=[],
                        images=[original_image],
                        screenshot=None,
                    )]
                    logger.info(
                        "Created synthetic page for standalone image (no OCR text)"
                    )
                elif len(pages[0].images) == 0:
                    # OCR found text but no embedded images detected — add original
                    pages[0].images.append(original_image)
                    logger.info(
                        "Added original image to first page for standalone image document"
                    )
            # ──────────────────────────────────────────────────────────────────

            # Detect document type
            doc_type = detect_document_type(result)

            logger.info(f"Docling extraction complete: {len(pages)} pages")
            return pages, doc_type

        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except Exception as e:
                logger.warning(f"Failed to delete temp file {tmp_path}: {e}")

    except ImportError:
        raise HTTPException(
            status_code=503, detail="Docling not available - cannot extract this document type"
        )
    except Exception as e:
        logger.error(f"Docling extraction failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Docling extraction failed: {str(e)}")


async def extract_with_llamaindex(
    content: bytes, filename: str, content_type: str, opts: ExtractionOptions
) -> tuple[List[PageData], str]:
    """
    Extract plain text documents using LlamaIndex.
    Returns as single page - chunking happens downstream in page-processing-worker
    to avoid double chunking and maintain consistency with other formats.

    Args:
        content: Raw file bytes
        filename: Original filename
        content_type: MIME type (text/plain only)
        opts: Extraction options (not used)

    Returns:
        Tuple of (single page with full text, document_type)
    """
    if not LLAMAINDEX_AVAILABLE:
        raise HTTPException(
            status_code=503, detail="LlamaIndex not available - cannot extract text documents"
        )

    logger.info(f"Extracting with LlamaIndex: {filename} ({content_type})")

    # Save to temp file (LlamaIndex readers require file path)
    with tempfile.NamedTemporaryFile(delete=False, suffix=FilePath(filename).suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        # Load document using LlamaIndex
        documents = load_text_document(tmp_path, content_type)

        # Convert to single page (NO chunking - happens in page-processing-worker)
        pages = convert_text_to_single_page(documents, opts)

        # Document type
        doc_type = "plain"

        logger.info(f"LlamaIndex extraction complete: {len(pages)} chunks")
        return pages, doc_type

    finally:
        # Clean up temp file
        try:
            os.unlink(tmp_path)
        except Exception as e:
            logger.warning(f"Failed to delete temp file {tmp_path}: {e}")


# ─── API Endpoints ───────────────────────────────────────────────────────────


@app.post("/extract", response_model=ExtractionResult)
async def extract_document(
    file: UploadFile = File(..., description="Document file to extract"),
    options: str = Form(default="{}", description="JSON string of ExtractionOptions"),
    language: Optional[str] = Form(
        default=None, description="Explicit language override (ISO 639-1 code, e.g., 'en', 'zh')"
    ),
    detectLanguage: bool = Form(default=True, description="Enable automatic language detection"),
    author: Optional[str] = Form(
        default=None, description="Explicit document author (overrides extracted metadata)"
    ),
    createdDate: Optional[str] = Form(
        default=None,
        description="Explicit creation date (ISO 8601 format, overrides extracted metadata)",
    ),
    modifiedDate: Optional[str] = Form(
        default=None,
        description="Explicit modification date (ISO 8601 format, overrides extracted metadata)",
    ),
    title: Optional[str] = Form(
        default=None, description="Explicit document title (overrides extracted metadata)"
    ),
    subject: Optional[str] = Form(
        default=None, description="Explicit document subject (overrides extracted metadata)"
    ),
    keywords: Optional[str] = Form(
        default=None,
        description="Explicit document keywords (comma-separated, overrides extracted metadata)",
    ),
):
    """
    Extract structured content from a document using unified extraction service.

    Supports 12 formats:
    - Docling (11): PDF, DOCX, PPTX, HTML, Markdown, PNG, JPEG, JPG, TIFF, BMP, WEBP
    - LlamaIndex (1): TXT

    Note: Legacy Office formats (.doc, .ppt) are NOT supported (require LibreOffice conversion).

    Returns:
    - Pages with text, layout, tables, images (unified format)
    - Document structure and metadata (including language detection and document properties)
    """
    start_time = datetime.now()
    content_type = "unknown"
    engine = "unknown"

    try:
        # Parse options
        import json

        opts_dict = json.loads(options)
        opts = ExtractionOptions(**opts_dict)

        # Get content type (prefer file.content_type, fallback to filename extension)
        content_type = file.content_type
        if not content_type or content_type == "application/octet-stream":
            # Guess from filename
            if file.filename.endswith(".txt"):
                content_type = "text/plain"
            elif file.filename.endswith(".md"):
                content_type = "text/markdown"
            elif file.filename.endswith(".json"):
                content_type = "application/json"
            elif file.filename.endswith(".csv"):
                content_type = "text/csv"
            elif file.filename.endswith(".xml"):
                content_type = "application/xml"

        logger.info(f"Starting extraction for file: {file.filename}, type: {content_type}")
        logger.info(f"Options: {opts.model_dump()}")

        # Read file content
        content = await file.read()
        file_size = len(content)
        logger.info(f"File size: {file_size} bytes")

        # Determine engine before routing
        if content_type in DOCLING_TYPES:
            engine = "docling"
        elif content_type in LLAMAINDEX_TYPES:
            engine = "llamaindex"
        else:
            engine = "unsupported"

        # Emit routing metric
        logger.info(
            f"[METRIC] extraction.route content_type={content_type} engine={engine} "
            f"file_size={file_size} filename={file.filename}"
        )

        # Route to appropriate extraction engine based on content type
        engine_start_time = datetime.now()

        if content_type in DOCLING_TYPES:
            # Docling path (PDF, Office docs, images, markdown, CSV)
            logger.info(f"Routing to Docling engine for {content_type}")
            pages, doc_type = await extract_with_docling(content, file.filename, opts)

        elif content_type in LLAMAINDEX_TYPES:
            # LlamaIndex path (plain text, JSON, XML)
            logger.info(f"Routing to LlamaIndex engine for {content_type}")
            pages, doc_type = await extract_with_llamaindex(
                content, file.filename, content_type, opts
            )

        else:
            # Unsupported format
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported content type: {content_type}. "
                f"Supported: {', '.join(sorted(DOCLING_TYPES | LLAMAINDEX_TYPES))}",
            )

        # Calculate metrics
        engine_duration = (datetime.now() - engine_start_time).total_seconds()
        processing_time = (datetime.now() - start_time).total_seconds()
        page_count = len(pages)
        table_count = sum(len(p.tables) for p in pages)
        image_count = sum(len(p.images) for p in pages)

        # Language detection
        language_info = None
        if language:
            # Explicit language override
            language_info = {
                "primary": language,
                "confidence": 1.0,
                "script": None,
                "method": "explicit",
                "secondary": [],
            }
            logger.info(f"Using explicit language override: {language}")
        elif detectLanguage and LANGUAGE_DETECTION_AVAILABLE and pages:
            # Auto-detect language from first few pages (performance optimization)
            try:
                # Sample from first 5 pages (or all pages if fewer)
                sample_pages = pages[: min(5, len(pages))]
                sample_text = "\n\n".join(p.text for p in sample_pages if p.text)

                if len(sample_text) >= 50:  # Minimum text length for detection
                    detector = get_language_detector()
                    language_info = detector.detect_with_cache(sample_text)
                    logger.info(
                        f"Language detected: {language_info['primary']} "
                        f"(confidence: {language_info['confidence']:.3f}, "
                        f"method: {language_info['method']})"
                    )
                else:
                    logger.info("Insufficient text for language detection")
            except Exception as e:
                logger.warning(f"Language detection failed: {e}")

        # Document metadata extraction
        doc_metadata = extract_document_metadata(content, file.filename, content_type)

        # Apply explicit metadata overrides (including empty strings)
        if author is not None:
            doc_metadata["author"] = author
        if title is not None:
            doc_metadata["title"] = title
        if subject is not None:
            doc_metadata["subject"] = subject
        if createdDate is not None:
            doc_metadata["createdDate"] = createdDate
        if modifiedDate is not None:
            doc_metadata["modifiedDate"] = modifiedDate
        if keywords is not None:
            # Parse keywords, empty string becomes empty list
            doc_metadata["keywords"] = (
                [k.strip() for k in keywords.split(",") if k.strip()] if keywords else []
            )

        logger.info(
            f"Document metadata: author={doc_metadata.get('author')}, "
            f"title={doc_metadata.get('title')}, "
            f"created={doc_metadata.get('createdDate')}"
        )

        # Build metadata
        metadata = ExtractionMetadata(
            pageCount=page_count,
            hasOCR=False,  # Set by Docling if applicable
            totalTables=table_count,
            totalImages=image_count,
            processingTime=processing_time,
            documentType=doc_type,
            # Language fields
            language=language_info["primary"] if language_info else None,
            languageConfidence=language_info["confidence"] if language_info else None,
            languageScript=language_info.get("script") if language_info else None,
            languageDetectionMethod=language_info.get("method") if language_info else None,
            secondaryLanguages=language_info.get("secondary") if language_info else None,
            # Document metadata fields
            author=doc_metadata.get("author"),
            createdDate=doc_metadata.get("createdDate"),
            modifiedDate=doc_metadata.get("modifiedDate"),
            title=doc_metadata.get("title"),
            subject=doc_metadata.get("subject"),
            keywords=doc_metadata.get("keywords"),
        )

        # Build structure
        structure = DocumentStructure(outline=[], documentType=doc_type)

        # Build result
        extraction_result = ExtractionResult(pages=pages, metadata=metadata, structure=structure)

        # Emit success metrics
        lang_detected = metadata.language or "none"
        lang_conf = metadata.languageConfidence or 0.0
        lang_method = metadata.languageDetectionMethod or "none"

        logger.info(f"Extraction complete in {processing_time:.2f}s - {page_count} pages extracted")
        logger.info(
            f"[METRIC] extraction.success content_type={content_type} engine={engine} "
            f"duration_total={processing_time:.3f} duration_engine={engine_duration:.3f} "
            f"pages={page_count} tables={table_count} images={image_count} "
            f"file_size={file_size} doc_type={doc_type} "
            f"language={lang_detected} language_confidence={lang_conf:.3f} language_method={lang_method}"
        )

        return extraction_result

    except HTTPException as e:
        # HTTP exceptions (like unsupported format) - emit metric
        error_type = "unsupported_format" if e.status_code == 400 else "http_error"
        processing_time = (datetime.now() - start_time).total_seconds()
        logger.info(
            f"[METRIC] extraction.error content_type={content_type} engine={engine} "
            f"error_type={error_type} status_code={e.status_code} duration={processing_time:.3f}"
        )
        raise
    except Exception as e:
        # Unexpected errors - emit metric
        processing_time = (datetime.now() - start_time).total_seconds()
        error_type = type(e).__name__
        logger.error(f"Extraction failed: {str(e)}", exc_info=True)
        logger.info(
            f"[METRIC] extraction.error content_type={content_type} engine={engine} "
            f"error_type={error_type} duration={processing_time:.3f}"
        )
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    try:
        # Check if Docling is available
        from docling.document_converter import DocumentConverter

        docling_available = True
    except ImportError:
        docling_available = False

    # Check LlamaIndex availability (set during import)
    llamaindex_available = LLAMAINDEX_AVAILABLE

    # Check language detection availability
    language_detection_available = LANGUAGE_DETECTION_AVAILABLE

    # Service is healthy if either engine is available
    is_healthy = docling_available or llamaindex_available

    status_code = 200 if is_healthy else 503
    response = {
        "status": "healthy" if is_healthy else "unhealthy",
        "service": "unified-extraction-service",
        "version": "2.2.0",  # Bumped for GPU acceleration support
        "timestamp": datetime.now().isoformat(),
        "engines": {
            "docling": {
                "available": docling_available,
                "formats": len(DOCLING_TYPES) if docling_available else 0,
            },
            "llamaindex": {
                "available": llamaindex_available,
                "formats": len(LLAMAINDEX_TYPES) if llamaindex_available else 0,
            },
        },
        "features": {
            "language_detection": language_detection_available,
            "gpu_acceleration": DEVICE == "cuda",
        },
        "compute": {
            "device": DEVICE,
            "gpu_name": GPU_DEVICE_NAME,
            "workers": int(os.environ.get("WORKERS", "2")),
        },
        "total_formats": (len(DOCLING_TYPES) if docling_available else 0)
        + (len(LLAMAINDEX_TYPES) if llamaindex_available else 0),
    }

    if not is_healthy:
        response["error"] = "Neither Docling nor LlamaIndex available"

    return JSONResponse(status_code=status_code, content=response) if not is_healthy else response


# ─── Main ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8080, reload=True, log_level="info")
