"""Explore Docling's document attribute for DOCX"""

import io
from docx import Document
from docling.document_converter import DocumentConverter
import tempfile
import os

# Create a DOCX
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is paragraph 1.')
doc.add_paragraph('This is paragraph 2.')
doc.add_heading('Section 2', 1)
doc.add_paragraph('More content here.')

buffer = io.BytesIO()
doc.save(buffer)
buffer.seek(0)
file_bytes = buffer.read()

# Save to temp file
with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
    tmp.write(file_bytes)
    tmp_path = tmp.name

try:
    converter = DocumentConverter()
    result = converter.convert(tmp_path)

    print(f"Pages: {len(result.pages)}")
    print(f"Document exists: {hasattr(result, 'document')}\n")

    if hasattr(result, 'document') and result.document:
        doc = result.document
        print(f"Document type: {type(doc)}")
        print(f"Document attributes: {[a for a in dir(doc) if not a.startswith('_')]}\n")

        # Try to export to markdown
        if hasattr(doc, 'export_to_markdown'):
            print("=== Document as Markdown ===")
            print(doc.export_to_markdown())
            print()

        if hasattr(doc, 'export_to_text'):
            print("=== Document as Text ===")
            print(doc.export_to_text())
            print()

        if hasattr(doc, 'main_text'):
            print(f"Main text: {doc.main_text[:200]}")

finally:
    if os.path.exists(tmp_path):
        os.unlink(tmp_path)
