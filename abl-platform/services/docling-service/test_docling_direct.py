"""Test Docling directly to understand format support"""

import io
from docx import Document
from docling.document_converter import DocumentConverter
import tempfile
import os

# Create a DOCX
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is paragraph 1.')
doc.add_paragraph('This is paragraph 2.')

buffer = io.BytesIO()
doc.save(buffer)
buffer.seek(0)
file_bytes = buffer.read()

print(f"Created DOCX: {len(file_bytes)} bytes\n")

# Save to temp file
with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
    tmp.write(file_bytes)
    tmp_path = tmp.name

print(f"Saved to: {tmp_path}\n")

try:
    # Try converting with Docling
    print("Initializing DocumentConverter...")
    converter = DocumentConverter()

    print(f"Converting {tmp_path}...\n")
    result = converter.convert(tmp_path)

    print(f"Result type: {type(result)}")
    print(f"Result attributes: {[a for a in dir(result) if not a.startswith('_')]}")

    if hasattr(result, 'pages'):
        print(f"\nNumber of pages: {len(result.pages)}")
        if len(result.pages) > 0:
            print(f"First page type: {type(result.pages[0])}")
            print(f"First page attributes: {[a for a in dir(result.pages[0]) if not a.startswith('_')]}")
    else:
        print("\nNo 'pages' attribute found!")

    if hasattr(result, 'document'):
        print(f"\nDocument attribute exists")
        print(f"Document type: {type(result.document)}")

    if hasattr(result, 'metadata'):
        print(f"\nMetadata: {result.metadata}")

finally:
    if os.path.exists(tmp_path):
        os.unlink(tmp_path)
