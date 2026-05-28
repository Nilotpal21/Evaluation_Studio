"""Check if DoclingDocument has page structure"""

import io
from docx import Document
from docling.document_converter import DocumentConverter
import tempfile
import os

# Create a DOCX
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is paragraph 1.')
doc.add_paragraph('This is paragraph 2.')

buffer = io.BytesIO()
doc.save(buffer)
buffer.seek(0)
file_bytes = buffer.read()

with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
    tmp.write(file_bytes)
    tmp_path = tmp.name

try:
    converter = DocumentConverter()
    result = converter.convert(tmp_path)

    doc = result.document

    print(f"num_pages: {doc.num_pages}")
    print(f"Has pages attr: {hasattr(doc, 'pages')}")

    if hasattr(doc, 'pages'):
        print(f"Pages type: {type(doc.pages)}")
        print(f"Pages length: {len(list(doc.pages))}")

        for i, page in enumerate(doc.pages):
            print(f"\nPage {i}:")
            print(f"  Type: {type(page)}")
            print(f"  Attributes: {[a for a in dir(page) if not a.startswith('_')][:10]}")

    # Check tables
    if hasattr(doc, 'tables'):
        tables_list = list(doc.tables)
        print(f"\nTables: {len(tables_list)}")

    # Check pictures
    if hasattr(doc, 'pictures'):
        pictures_list = list(doc.pictures)
        print(f"Pictures: {len(pictures_list)}")

finally:
    if os.path.exists(tmp_path):
        os.unlink(tmp_path)
